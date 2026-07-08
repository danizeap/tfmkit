"""PUBLISHER — assembles a Markdown document into a formatted .docx. NO LLM, no network.

Every cover datum (university, degree, title, author, tutors, academic year) and every
typography value comes from the per-project tfm.config.yaml. Nothing is hardcoded: a missing
value renders as a visible [PENDIENTE: …] placeholder so it can never silently ship wrong.
The AI-usage declaration is a generalized, honest template; the tool name comes from config.

Usage:
  python scripts/publish.py --in <document.md> --config tfm.config.yaml [--out <file.docx>]

Run scripts/lint.py first: publishing with unresolved markers is the author's decision,
not the tool's.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")
    except Exception:
        pass

DECLARACION = """## Anexo — Declaración de uso de IA generativa

En la elaboración de este documento se ha empleado una herramienta de IA generativa como apoyo en
las siguientes fases: {fases}. El contenido sustantivo, los datos, los resultados y las
conclusiones son obra del autor. La herramienta ha trabajado exclusivamente a partir de los
materiales y respuestas aportados por el autor, y no ha generado resultados ni referencias
inexistentes: todas las referencias citadas han sido verificadas.

- Herramienta / modelo utilizado: {herramienta}
- Fases asistidas: {fases}.
- El autor asume la autoría y la responsabilidad plena del contenido.
"""


def load_config(config_path: str | None) -> dict:
    if not config_path:
        return {}
    try:
        import yaml
    except ImportError:
        raise SystemExit("Falta PyYAML para leer la configuración. Instala: pip install pyyaml")
    return yaml.safe_load(Path(config_path).read_text(encoding="utf-8")) or {}


def cover_value(cfg_project: dict, key: str, label: str) -> str:
    return str(cfg_project.get(key) or "").strip() or f"[PENDIENTE: {label}]"


def strip_fact_tags(md: str) -> str:
    return re.sub(r"\s*\[F\d+\]", "", md)


def _field(paragraph, instr: str, placeholder: str = "") -> None:
    """Insert a Word field (PAGE, TOC…). Word computes it; for TOC the reader updates with F9."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run._r.append(begin)
    run._r.append(instr_el)
    if placeholder:
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        run._r.append(sep)
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def build_docx(md: str, out: str, cfg: dict) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    project = cfg.get("project") or {}
    typo = cfg.get("typography") or {}

    d = docx.Document()

    style = d.styles["Normal"]
    if typo.get("font"):
        style.font.name = str(typo["font"])
    if typo.get("size_pt"):
        style.font.size = Pt(float(typo["size_pt"]))
    if typo.get("line_spacing"):
        style.paragraph_format.line_spacing = float(typo["line_spacing"])

    # Headings follow the configured font too (Word's theme default is a different,
    # colored font — an instant guideline violation if left as-is).
    base_pt = float(typo.get("size_pt") or 12)
    for name, delta in (("Heading 1", 4), ("Heading 2", 2), ("Heading 3", 1), ("Title", 8)):
        h = d.styles[name]
        if typo.get("font"):
            h.font.name = str(typo["font"])
        h.font.size = Pt(base_pt + delta)
        h.font.color.rgb = RGBColor(0, 0, 0)

    margins = typo.get("margins_cm") or {}
    for section in d.sections:
        if margins.get("top"):
            section.top_margin = Cm(float(margins["top"]))
        if margins.get("bottom"):
            section.bottom_margin = Cm(float(margins["bottom"]))
        if margins.get("left"):
            section.left_margin = Cm(float(margins["left"]))
        if margins.get("right"):
            section.right_margin = Cm(float(margins["right"]))

    # Page number in the footer of every page except the cover.
    section = d.sections[0]
    section.different_first_page_header_footer = True
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(footer_p, "PAGE")

    # Cover — config values only; a hole is a visible [PENDIENTE], never a default.
    d.add_paragraph(cover_value(project, "university", "universidad"))
    d.add_paragraph(cover_value(project, "degree", "titulación"))
    p = d.add_paragraph()
    run = p.add_run(cover_value(project, "title", "título del TFM"))
    run.bold = True
    run.font.size = Pt(22)
    d.add_paragraph(cover_value(project, "author", "autor"))
    for tutor in project.get("tutors") or []:
        role = str((tutor or {}).get("role") or "").strip() or "[PENDIENTE: rol del tutor]"
        name = str((tutor or {}).get("name") or "").strip() or "[PENDIENTE: nombre del tutor]"
        d.add_paragraph(f"{role}: {name}")
    d.add_paragraph("Curso académico: " + cover_value(project, "academic_year", "curso académico"))
    d.add_page_break()

    # Paginated table of contents (Word recomputes it: Ctrl+A then F9, or on print).
    if typo.get("include_toc", True):
        toc_title = d.add_paragraph()
        toc_run = toc_title.add_run("Índice")
        toc_run.bold = True
        toc_run.font.size = Pt(base_pt + 4)
        _field(d.add_paragraph(), r'TOC \o "1-3" \h \z \u',
               "Índice pendiente de actualizar: seleccione todo (Ctrl+A) y pulse F9.")
        d.add_page_break()

    for line in md.splitlines():
        s = line.rstrip()
        if s.startswith("### "):
            d.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            d.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            d.add_heading(s[2:], level=1)
        elif s.strip():
            d.add_paragraph(s)
    d.save(out)


def main():
    ap = argparse.ArgumentParser(description="Publisher: Markdown -> .docx con portada desde config.")
    ap.add_argument("--in", dest="infile", required=True, help="documento .md")
    ap.add_argument("--out", help="salida .docx (por defecto <in>.docx)")
    ap.add_argument("--config", help="tfm.config.yaml del proyecto")
    args = ap.parse_args()

    md = Path(args.infile).read_text(encoding="utf-8")
    cfg = load_config(args.config)

    body = strip_fact_tags(md)
    disclosure = cfg.get("ai_disclosure") or {}
    if disclosure.get("include", True):
        herramienta = str(disclosure.get("tool") or "").strip() or "[PENDIENTE: herramienta y versión]"
        fases = str(disclosure.get("assisted_phases") or "").strip() or "[PENDIENTE: fases asistidas]"
        body += "\n\n" + DECLARACION.format(herramienta=herramienta, fases=fases)

    out = args.out or (str(args.infile).rsplit(".md", 1)[0] + ".docx")
    try:
        build_docx(body, out, cfg)
    except ImportError as exc:
        raise SystemExit(f"Falta python-docx ({exc}). Instala: pip install python-docx")
    print(f"✔ Documento generado: {out}")


if __name__ == "__main__":
    main()
